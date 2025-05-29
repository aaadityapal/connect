from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

# Check if logo file exists, if not create placeholder message
logo_path = "lol.jpeg"
logo_exists = os.path.exists(logo_path)

# Create Word Document with updated details
doc = Document()

# Add logo to the header
section = doc.sections[0]
header = section.header
header_para = header.paragraphs[0]
header_para.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

if logo_exists:
    run = header_para.add_run()
    run.add_picture(logo_path, width=Inches(1.5))  # Adjust width as needed
else:
    header_para.add_run("Logo placeholder (lol.jpeg not found)")

# Add footer with contact information
footer = section.footer

# Add page number to the top left side
footer_para = footer.add_paragraph()
footer_para.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
footer_para.paragraph_format.space_after = Pt(0)
page_num = footer_para.add_run("P a g e  ")
page_num.font.size = Pt(8)

# Add field for page number
run = footer_para.add_run()
fldChar = OxmlElement('w:fldChar')
fldChar.set(qn('w:fldCharType'), 'begin')
run._element.append(fldChar)

instrText = OxmlElement('w:instrText')
instrText.set(qn('xml:space'), 'preserve')
instrText.text = 'PAGE'
run._element.append(instrText)

fldChar = OxmlElement('w:fldChar')
fldChar.set(qn('w:fldCharType'), 'end')
run._element.append(fldChar)
run.font.size = Pt(8)

# Add horizontal line below page number
footer_para = footer.add_paragraph()
p = footer_para._p  # Get the paragraph element
pPr = p.get_or_add_pPr()  # Get or create paragraph properties
pBdr = OxmlElement('w:pBdr')  # Create paragraph border element
pPr.append(pBdr)  # Add border element to paragraph properties

# Add top border (horizontal line)
top = OxmlElement('w:top')
top.set(qn('w:val'), 'single')
top.set(qn('w:sz'), '6')  # Border width
top.set(qn('w:space'), '1')
top.set(qn('w:color'), 'auto')
pBdr.append(top)

# Add contact information right-justified in the footer
# First line
footer_para = footer.add_paragraph()
footer_para.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
footer_para.paragraph_format.space_before = Pt(0)
contact_line1 = footer_para.add_run("H.O.: F-52, First Floor, Lane No.16, Madhu Vihar,")
contact_line1.font.size = Pt(8)

# Second line
footer_para = footer.add_paragraph()
footer_para.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
footer_para.paragraph_format.space_before = Pt(0)
footer_para.paragraph_format.space_after = Pt(0)
contact_line2 = footer_para.add_run("I. P. Extension, Delhi-110092 Phone: 9958600397, 7503477154")
contact_line2.font.size = Pt(8)

# Third line with email and website as hyperlinks
footer_para = footer.add_paragraph()
footer_para.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
footer_para.paragraph_format.space_before = Pt(0)
footer_para.add_run("E-Mail:").font.size = Pt(8)

# Add email as hyperlink
hyperlink_id = f"email_link"
email_address = "contact@architectshive.com"
email_run = footer_para.add_run(" ")
email_run.font.size = Pt(8)

# Create hyperlink for email
hyperlink_rel = footer_para.part.relate_to(f"mailto:{email_address}", "hyperlink", is_external=True)
hyperlink = OxmlElement('w:hyperlink')
hyperlink.set(qn('r:id'), hyperlink_rel)
hyperlink.set(qn('w:history'), '1')

# Create properties element
rPr = OxmlElement('w:rPr')
# Add color
c = OxmlElement('w:color')
c.set(qn('w:val'), '0000FF')  # Blue color
rPr.append(c)
# Add underline
u = OxmlElement('w:u')
u.set(qn('w:val'), 'single')
rPr.append(u)

# Create run element
r = OxmlElement('w:r')
r.append(rPr)
t = OxmlElement('w:t')
t.text = email_address
r.append(t)
hyperlink.append(r)

# Add hyperlink to run
email_run._r.append(hyperlink)

footer_para.add_run(", Website:").font.size = Pt(8)

# Add website as hyperlink
website_run = footer_para.add_run(" ")
website_run.font.size = Pt(8)
website_address = "www.architectshive.com"

# Create hyperlink for website
hyperlink_rel = footer_para.part.relate_to(f"http://{website_address}", "hyperlink", is_external=True)
hyperlink = OxmlElement('w:hyperlink')
hyperlink.set(qn('r:id'), hyperlink_rel)
hyperlink.set(qn('w:history'), '1')

# Create properties element
rPr = OxmlElement('w:rPr')
# Add color
c = OxmlElement('w:color')
c.set(qn('w:val'), '0000FF')  # Blue color
rPr.append(c)
# Add underline
u = OxmlElement('w:u')
u.set(qn('w:val'), 'single')
rPr.append(u)

# Create run element
r = OxmlElement('w:r')
r.append(rPr)
t = OxmlElement('w:t')
t.text = website_address
r.append(t)
hyperlink.append(r)

# Add hyperlink to run
website_run._r.append(hyperlink)

doc.add_heading('Minutes of Meeting (MoM) – Construction Site Visit', 0)

# Meeting Details
doc.add_paragraph('Date: 25/05/2025')
doc.add_paragraph('Time: Afternoon (2:30 PM – 4:00 PM)')
doc.add_paragraph('Location: Flat Number 4041, Jasola, New Delhi')
doc.add_paragraph('Project Name: 4041, Jasola, New Delhi')
doc.add_paragraph('Meeting Type: Progress Review and Coordination')
doc.add_paragraph('Meeting Conducted By: Mr. Khwaza Fratulla (Client), Mr. Faraz Fratulla (Client), Rachana Pal & Prabhat Arya (Architects), Tanis Shil (Site Coordinator) ')

# Attendees
doc.add_heading('Attendees', level=1)
table = doc.add_table(rows=1, cols=4)
table.style = 'Table Grid'  # Add visible grid lines

# Make header row bold
hdr_cells = table.rows[0].cells
for cell in hdr_cells:
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.bold = True
    
hdr_cells[0].text = 'Name'
hdr_cells[1].text = 'Designation'
hdr_cells[2].text = 'Organization/Company'
hdr_cells[3].text = 'Contact Info'

# Make all header text bold again (since setting text removes formatting)
for cell in hdr_cells:
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.bold = True

attendees = [
    ('Mr. Khwaza Fratulla', 'Client', '-', '-'),
    ('Mr. Faraz Fratulla', 'Client', '-', '-'),
    ('Prabhat Arya', 'Architect', 'ArchitectsHive', '+91 7503468992'),
    ('Rachana Pal', 'Architect', 'ArchitectsHive', '+91 7503468992'),
    ('Tanis Shil', 'Site Coordinator', 'ArchitectsHive', '+91 9211549809'),


]

for name, desig, org, contact in attendees:
    row_cells = table.add_row().cells
    row_cells[0].text = name
    row_cells[1].text = desig
    row_cells[2].text = org
    row_cells[3].text = contact

# Agenda
doc.add_heading('Agenda', level=1)
agenda_items = [
    'Drawings discussion and confirmation of accuracy',
    'Completion of pending works at the earliest',
    'Strategies to accelerate site progress',
    'Final approvals and clarifications'
]
for item in agenda_items:
    doc.add_paragraph(item, style='List Bullet')

# Key Discussions & Decisions
doc.add_heading('Key Discussions & Decisions', level=1)
discussions = [
    "All drawings reviewed and confirmed as accurate by the architects and client.",
    "Site coordinator to prioritize and finish the pending works on high priority.",
    "Team discussed measures to improve workflow and manpower for faster site completion.",
    "Final approvals were given on minor layout changes and clarifications addressed."
]
for point in discussions:
    doc.add_paragraph(f'- {point}')

# Site Observations
doc.add_heading('Site Observations', level=1)
observations = [
    'Overall construction progress is satisfactory.',
    'Electrical works and carpentry are ongoing.',
    'Site is maintained well with adequate safety protocols.'
]
for item in observations:
    doc.add_paragraph(item, style='List Bullet')

# Client Approvals/Instructions
doc.add_heading('Client Approvals/Instructions', level=1)
approvals = [
    'Approved layout as per latest drawing revision.',
    'Requested completion of A.C. line works and revised electrical point markings.'
]
for item in approvals:
    doc.add_paragraph(item, style='List Bullet')

# Concerns Raised
doc.add_heading('Concerns Raised', level=1)
concerns = [
    "Changes in the electrical points between site and drawings.",
    "A.C. pipe cutting and wall coordination to be completed by Tuesday.",
    "Final selection of lighting fixtures required.",
    "One-week work schedule to be submitted to the client.",
    "Updated revised drawings to be shared as per today's discussion.",
    "Plumbing work to be initiated this week.",
    "Bathroom waterproofing details and execution plan needed."
]
for point in concerns:
    doc.add_paragraph(f'- {point}')

# Next Steps
doc.add_heading('Next Steps', level=1)
next_steps = [
    'Submit revised electrical drawings within 2 days.',
    'Complete pending wall cuttings for A.C. pipes by Tuesday.',
    'Prepare and submit one-week work schedule.',
    'Initiate plumbing and waterproofing works as planned.'
]
for step in next_steps:
    doc.add_paragraph(step, style='List Bullet')


# Signatures - Updated to match the provided format
doc.add_heading('Signatures', level=1)
doc.add_paragraph('Name | Designation | Signature | Date')

# Add signature lines for each attendee
signatures = [
    ('Mr. Khwaza Fratulla', 'Client', '25/05/2025'),
    ('Tanis Shil', 'Site Coordinator', '25/05/2025'),
    ('Rachana Pal', 'Architect', '25/05/2025'),
    ('Prabhat Arya', 'Architect', '25/05/2025')
]

for name, designation, date in signatures:
    p = doc.add_paragraph()
    p.add_run(f"{name} | {designation} | ").bold = False
    p.add_run("____________").bold = False
    p.add_run(f" | {date}").bold = False
    doc.add_paragraph()  # Add empty line between signatures

# Save the Word document
formatted_word_path = "MoM_Construction_Site_Formatted_v2.docx"
doc.save(formatted_word_path)

print(f"Word document saved as: {formatted_word_path}")
